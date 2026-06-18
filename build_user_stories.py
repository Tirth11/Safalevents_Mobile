# -*- coding: utf-8 -*-
"""Generate the End-to-End User Stories Word document for the SafalEvents MOBILE app."""
import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = r"E:\Safalvir\SafalEvents-Mobile\SafalEvents_Mobile_User_Stories.docx"
TODAY = datetime.date(2026, 6, 18).strftime("%d %B %Y")

BRAND = RGBColor(0xF2, 0x54, 0x1B)      # SafalEvents orange
INK   = RGBColor(0x1F, 0x29, 0x37)
MUTED = RGBColor(0x6B, 0x72, 0x80)
LABELBG = "FBEEE8"
HEADBG  = "F2541B"

# ----------------------------------------------------------------------------
# MOBILE USER STORY DATA  (Expo / React Native app, grounded in the codebase)
# ----------------------------------------------------------------------------
EPICS = [
 {"id":"NAV","name":"App Shell, Navigation & Platform","desc":"The mobile app foundation: splash, role-based navigation, the phone-frame web rendering, theming and the in-memory live store.","stories":[
   {"id":"US-NAV-001","title":"Splash & first launch routing","role":"user opening the app","want":"see a branded splash that then takes me into the right place","benefit":"the app starts fast and lands me correctly","pri":"High","deps":"—",
    "ac":["A branded splash shows briefly on launch.","A remembered Guest/Host session is routed straight to its flow.","Otherwise the app opens into Guest Mode (browse) with no login wall."]},
   {"id":"US-NAV-002","title":"Role-based bottom-tab navigation","role":"user","want":"navigate via bottom tabs appropriate to my context","benefit":"I can move around the app easily","pri":"High","deps":"US-NAV-001",
    "ac":["Guest Mode tabs: Discover, Explore, Account.","Guest tabs: Tickets, Explore, Messages, Profile.","Host tabs: Dashboard, Events, Messages, Account.","Staff tabs are limited to the role's permitted modules."]},
   {"id":"US-NAV-003","title":"Phone-frame rendering on web","role":"reviewer on a browser","want":"the app to look like a handset when opened on the web","benefit":"I can preview the mobile UI without a device","pri":"Medium","deps":"—",
    "ac":["On web the app renders inside a phone-frame device mockup.","A reserved status bar keeps content from being clipped under the notch.","On a real device the app renders full-screen."]},
   {"id":"US-NAV-004","title":"Consistent theming & components","role":"user","want":"a consistent, branded look across screens","benefit":"the app feels polished and easy to use","pri":"Medium","deps":"—",
    "ac":["Shared theme (SafalEvents orange/green) and a reusable component kit are used throughout.","Cards, badges, buttons, tabs and inputs are visually consistent.","Status/approval states use consistent colour coding."]},
   {"id":"US-NAV-005","title":"In-memory live store","role":"user","want":"actions like check-in to reflect immediately across screens","benefit":"the prototype behaves like a connected app","pri":"Medium","deps":"—",
    "ac":["A lightweight store lets a change (e.g. a check-in) notify subscribed screens.","The host dashboard reflects staff actions without a reload.","No real backend is required (prototype data)."]},
 ]},

 {"id":"BROWSE","name":"Guest Mode — Browse before login","desc":"Letting a new user explore the app freely and only hit auth when an action needs an account (UC-14).","stories":[
   {"id":"US-BROWSE-001","title":"Open into Guest Mode (no login wall)","role":"newly installed user","want":"explore the app without logging in or registering","benefit":"I get a feel for it before committing","pri":"High","deps":"US-NAV-001",
    "ac":["After the splash the app opens into a browse home with no login wall.","I can scroll and explore without an account.","Guest Mode is read-only - nothing writes until I authenticate."]},
   {"id":"US-BROWSE-002","title":"Discover landing","role":"guest-mode user","want":"a landing screen that showcases events and categories","benefit":"I quickly see what's on","pri":"Medium","deps":"US-BROWSE-001",
    "ac":["A hero section communicates 'browse freely, no account needed'.","Category chips and a trending list are shown.","Selecting an event opens its detail screen."]},
   {"id":"US-BROWSE-003","title":"Explore & search events","role":"guest-mode user","want":"browse and search the full event list","benefit":"I can find something specific","pri":"Medium","deps":"US-BROWSE-001",
    "ac":["An Explore tab lists events with cover, title, rating and meta.","A search field and category chips narrow the list.","Cards open the event detail screen."]},
   {"id":"US-BROWSE-004","title":"Open event detail without an account","role":"guest-mode user","want":"see full event details before deciding","benefit":"I have what I need to decide","pri":"High","deps":"US-BROWSE-001",
    "ac":["The detail screen shows cover, schedule, venue, host and description.","RSVP and Message entry points are visible.","No account is required to view."]},
   {"id":"US-BROWSE-005","title":"Persistent Login / Sign up entry","role":"guest-mode user","want":"an always-visible way to sign in","benefit":"I can create an account whenever I'm ready","pri":"Medium","deps":"US-BROWSE-001",
    "ac":["A Login / Sign up entry is visible while browsing (header + Account tab).","It opens the Auth screen.","If already signed in it shows 'My account' instead."]},
   {"id":"US-BROWSE-006","title":"Gated actions trigger auth","role":"guest-mode user","want":"to be sent to Login/Register only when an action needs an account","benefit":"I'm never blocked until necessary","pri":"High","deps":"US-BROWSE-001, US-AUTH-001",
    "ac":["RSVP, Message host, Save, Create event and My Dashboard route to the Auth screen.","The triggering action (intent) is remembered.","All other browsing remains open."]},
   {"id":"US-BROWSE-007","title":"Return to intent after auth","role":"guest-mode user","want":"to be returned to the exact action I tapped after signing in","benefit":"the flow feels seamless","pri":"High","deps":"US-BROWSE-006",
    "ac":["After login/registration the app opens the screen/action that triggered auth (e.g. the RSVP).","Selected-event context is preserved.","Routing falls back to the role home if no intent is stored."]},
   {"id":"US-BROWSE-008","title":"Cancel auth returns to Guest Mode","role":"guest-mode user","want":"to dismiss the auth screen and keep browsing","benefit":"I'm never trapped in a login wall","pri":"Medium","deps":"US-BROWSE-006",
    "ac":["The Auth screen is dismissible (close / back / swipe).","Dismissing returns to Guest Mode with nothing changed.","No partial account is created."]},
   {"id":"US-BROWSE-009","title":"Read-only Guest Mode","role":"platform owner","want":"no data to be written until the user authenticates","benefit":"guest browsing is safe and side-effect free","pri":"Medium","deps":"US-BROWSE-001",
    "ac":["Browsing performs no writes.","Save/RSVP/message all require auth first.","State changes begin only after sign-in."]},
   {"id":"US-BROWSE-010","title":"Remembered session & deep links","role":"returning or linked user","want":"to skip Guest Mode when remembered, and have links respect gating","benefit":"returning and shared-link journeys are smooth","pri":"Low","deps":"US-BROWSE-001",
    "ac":["The device remembers the session (web) so returning users skip Guest Mode.","Deep links / shared event links open the target in Guest Mode.","If the target needs auth, the same gated redirect applies, then returns to the target."]},
 ]},

 {"id":"AUTH","name":"Authentication, Signup & Login","desc":"Account creation and sign-in on mobile: Guest vs Host signup (Individual / Organization), OTP, unified login and the staff Invite-ID path.","stories":[
   {"id":"US-AUTH-001","title":"Choose Guest or Host signup","role":"new user","want":"clearly choose to register as a Guest or a Host","benefit":"I land in the right experience","pri":"High","deps":"—",
    "ac":["The Auth screen offers Sign up / Log in and a Guest / Host choice.","Selecting Host reveals an Individual / Organization choice.","Staff do not sign up here (they use an Invite ID)."]},
   {"id":"US-AUTH-002","title":"Host signup — Individual","role":"individual host","want":"register with my name and email/phone","benefit":"I can create and manage my own events","pri":"High","deps":"US-AUTH-001, US-AUTH-004",
    "ac":["First/last name and email/phone are captured.","An OTP step verifies the contact.","On success an individual host account is created and lands in the Host flow."]},
   {"id":"US-AUTH-003","title":"Host signup — Organization","role":"organization representative","want":"register my organization with its details and documents","benefit":"my org can host under a verified profile","pri":"High","deps":"US-AUTH-001, US-AUTH-004",
    "ac":["Org name, type (NGO/Temple/Company/Community/Other), website, city/state and contact person are captured.","An 'upload verification documents' action and an authorized-to-represent consent are present.","After OTP the app indicates the application is submitted for admin review."]},
   {"id":"US-AUTH-004","title":"OTP verification","role":"registering user","want":"confirm my email/phone with a one-time code","benefit":"only verified people create accounts","pri":"High","deps":"—",
    "ac":["A 6-digit code step follows the signup form.","The demo code is shown for testing; a clear error guides invalid entries.","Verification completes the account creation."]},
   {"id":"US-AUTH-005","title":"Login routes by account type","role":"registered Guest or Host","want":"log in with email/phone and reach the right flow","benefit":"one login does the right thing","pri":"High","deps":"US-AUTH-001",
    "ac":["Email/phone login signs the user in.","Host -> Host flow; Guest -> Guest flow.","A pending intent (from a gated action) takes priority for routing."]},
   {"id":"US-AUTH-006","title":"Login as Staff (Invite ID)","role":"invited staff member","want":"sign in with my Invite ID and email/phone","benefit":"I reach my assigned event with only my role's screens","pri":"High","deps":"US-AUTH-001, US-STAFF-001",
    "ac":["A 'Login as Staff' path collects Invite ID + email/phone.","The Invite ID is validated against the email it was issued to.","A wrong/expired/revoked invite is rejected with a clear message.","On success the staff lands in their event with role-limited tabs."]},
   {"id":"US-AUTH-007","title":"Sign out / switch account","role":"signed-in user","want":"sign out or switch to another account","benefit":"I can change context on a shared device","pri":"Medium","deps":"US-AUTH-005",
    "ac":["Profile/Account screens offer 'Switch account' and 'Log out'.","Signing out clears the session (and any staff context) and returns to Guest Mode / Auth.","Remembered session is cleared on log out."]},
 ]},

 {"id":"GUEST","name":"Guest — RSVP, Tickets & Messaging","desc":"The authenticated guest experience on mobile: RSVP flow, tickets with approval states, QR pass, messaging and profile.","stories":[
   {"id":"US-GUEST-001","title":"Complete an RSVP","role":"guest","want":"RSVP through a guided flow","benefit":"I can register for an event","pri":"High","deps":"US-BROWSE-004",
    "ac":["The flow captures contact details, party size and any host questions.","Paid events show ticket/payment fields.","Approval-required events explain the request will be reviewed; submit confirms."]},
   {"id":"US-GUEST-002","title":"See approval / waitlist states on my tickets","role":"guest","want":"clear status for each of my RSVPs","benefit":"I know exactly where I stand","pri":"High","deps":"US-GUEST-001",
    "ac":["My Tickets shows Under Approval / Waitlisted / Confirmed / Not Approved.","The action button reflects the state (View Pass vs disabled).","An 'action required' indicator highlights pending items."]},
   {"id":"US-GUEST-003","title":"View my ticket pass & QR","role":"guest","want":"open my QR pass for a confirmed event","benefit":"I can be checked in at the gate","pri":"High","deps":"US-GUEST-002",
    "ac":["A ticket-style pass shows event, booking id and a QR placeholder.","The pass is available only for approved/confirmed RSVPs.","Add-to-calendar and (if enabled) Message host are offered."]},
   {"id":"US-GUEST-004","title":"My Tickets list","role":"guest","want":"see all my RSVPs in one place","benefit":"I can track what I'm attending","pri":"High","deps":"US-GUEST-001",
    "ac":["Tickets are listed with cover, title, date, host and status.","Pending/waitlist items are surfaced.","A history/records section is shown."]},
   {"id":"US-GUEST-005","title":"Message the host (when enabled)","role":"guest","want":"message the host for an event that allows it","benefit":"I can ask questions","pri":"Medium","deps":"US-MSG-001",
    "ac":["The Message host option appears only when the event enables messaging.","A chat thread shows host/guest bubbles.","A composer lets me send (prototype)."]},
   {"id":"US-GUEST-006","title":"Guest messages inbox","role":"guest","want":"see my conversations and delivery logs","benefit":"I keep track of communications","pri":"Low","deps":"US-GUEST-005",
    "ac":["A Messages tab lists conversations and opens a thread.","A delivery-logs view lists messages sent to me.","Empty states are shown when there is nothing."]},
   {"id":"US-GUEST-007","title":"Guest profile & settings","role":"guest","want":"see my profile and notification preferences","benefit":"I control my experience","pri":"Low","deps":"US-AUTH-005",
    "ac":["Profile shows my identity and simple stats.","Notification preference toggles are present.","Switch account / Log out are available."]},
 ]},

 {"id":"HOST","name":"Host Tools","desc":"The host experience on mobile: dashboard, events, event creation and the per-event management hub.","stories":[
   {"id":"US-HOST-001","title":"Host dashboard","role":"host","want":"an overview of my events and activity","benefit":"I stay on top of things","pri":"High","deps":"US-AUTH-005",
    "ac":["Stat cards show events, RSVPs, pending approvals and check-in rate.","A 'RSVPs awaiting approval' highlight links to review.","A notifications entry and Create event action are present."]},
   {"id":"US-HOST-002","title":"Events list","role":"host","want":"see and open my events","benefit":"I can manage each one","pri":"High","deps":"US-HOST-001",
    "ac":["Events are listed with cover, status, date and RSVP count.","Tabs/search help find events.","A Manage action opens the event hub."]},
   {"id":"US-HOST-003","title":"Create an event","role":"host","want":"create an event with its key settings","benefit":"I can launch a new event","pri":"High","deps":"US-HOST-001",
    "ac":["Core fields (title, type, date, time, location, capacity, price) are captured.","Toggles include Require RSVP approval and Allow guest messaging.","Custom questions can be added; publishing returns to events."]},
   {"id":"US-HOST-004","title":"Manage an event (hub)","role":"host","want":"a tabbed hub to manage one event","benefit":"everything for an event is in one place","pri":"High","deps":"US-HOST-002",
    "ac":["Tabs: Overview, Guests, Messaging, Staff & Roles, Settings, Check-in.","Overview shows RSVP/going/pending/waitlist stats and recent activity.","Tabs respect the viewer's permissions."]},
   {"id":"US-HOST-005","title":"Event settings & notifications","role":"host","want":"edit an event's settings on mobile","benefit":"I can adjust an event on the go","pri":"Medium","deps":"US-HOST-004",
    "ac":["Settings edit core fields and toggles (approval, messaging, self-edit, paid).","A notifications screen lists activity (new RSVPs, messages, check-ins).","Changes are reflected in the relevant tabs."]},
   {"id":"US-HOST-006","title":"Host account","role":"host","want":"manage my account, earnings and settings","benefit":"I control my host profile","pri":"Low","deps":"US-AUTH-005",
    "ac":["Account shows profile, an earnings/payouts summary and settings toggles.","Switch account / Log out are available.","Host-only entries are hidden from staff."]},
 ]},

 {"id":"APPROVAL","name":"Approval & Waitlist (Mobile)","desc":"The Under-Approval lifecycle and capacity waitlist on mobile, reflected for host, guest and staff.","stories":[
   {"id":"US-APPROVAL-001","title":"Review the Under-Approval queue","role":"host or permitted staff","want":"approve or reject pending RSVPs","benefit":"I curate who attends","pri":"High","deps":"US-HOST-004",
    "ac":["Pending RSVPs are listed with response, party size and answers.","Approve and Reject actions are available to permitted roles.","Approving admits the guest; the list updates."]},
   {"id":"US-APPROVAL-002","title":"Waitlist held under approval","role":"host","want":"capacity-full RSVPs to wait for my approval","benefit":"I decide who gets a freed seat","pri":"High","deps":"US-APPROVAL-001",
    "ac":["When full, new RSVPs are waitlisted and held under approval.","Approving a waitlisted guest admits them.","The host is prompted to review the waitlist when a seat frees."]},
   {"id":"US-APPROVAL-003","title":"Approval status on guest screens","role":"guest","want":"see my approval state on my tickets and the event","benefit":"there's no ambiguity","pri":"Medium","deps":"US-GUEST-002",
    "ac":["Under Approval / Waitlisted / Approved / Not Approved render consistently.","The QR pass appears only when approved.","Badges use consistent colours."]},
   {"id":"US-APPROVAL-004","title":"Host dashboard reflects changes live","role":"host","want":"approvals and check-ins to update my dashboard immediately","benefit":"I see accurate numbers in real time","pri":"Medium","deps":"US-NAV-005",
    "ac":["A staff check-in updates the host dashboard live.","Approval decisions update counts and lists.","No manual refresh is needed."]},
 ]},

 {"id":"STAFF","name":"Staff, Roles & QR Check-in","desc":"The staff experience: invite-based login, permission-limited modules and gate check-in by scanning guest QR passes.","stories":[
   {"id":"US-STAFF-001","title":"Staff invite login & role-limited tabs","role":"staff member","want":"sign in via Invite ID and see only my permitted screens","benefit":"I do my job with the right access","pri":"High","deps":"US-AUTH-006",
    "ac":["Signing in with a valid Invite ID opens the staff portal.","Only role-permitted tabs are shown (e.g. Check-in only for a QR Scanner).","The staff member is scoped to their assigned event(s)."]},
   {"id":"US-STAFF-002","title":"Scan & validate a guest QR","role":"gate staff","want":"scan a guest's pass and validate it","benefit":"only valid, approved guests are admitted","pri":"High","deps":"US-STAFF-001, US-APPROVAL-001",
    "ac":["A scan screen lets me scan/enter a pass id.","Valid + approved + not-checked-in -> show guest details and allow check-in.","Already-checked-in -> warn with time; invalid/not-approved/wrong-event -> reject."]},
   {"id":"US-STAFF-003","title":"Mark arrived -> live update + email","role":"gate staff","want":"check a guest in and have it propagate","benefit":"attendance is tracked accurately","pri":"High","deps":"US-STAFF-002, US-NAV-005",
    "ac":["Marking arrived updates the live arrival count.","The host dashboard reflects the check-in immediately.","A confirmation email to the guest is logged."]},
   {"id":"US-STAFF-004","title":"Manage staff & roles (host)","role":"host","want":"invite staff and define roles on mobile","benefit":"I can run my team from my phone","pri":"Medium","deps":"US-HOST-004",
    "ac":["The Staff & Roles tab lists team members with status and role.","An invite form (name, email, role) is present; roles show their permission counts.","An example permission checklist explains module/action access."]},
   {"id":"US-STAFF-005","title":"Permission-gated modules (default-deny)","role":"staff member","want":"to be shown only what my role allows","benefit":"I cannot access unauthorized functions","pri":"High","deps":"US-STAFF-001",
    "ac":["Tabs and actions render based on role permissions.","Anything not granted is hidden (default-deny).","A profile screen shows my role and what it allows."]},
 ]},

 {"id":"MSG","name":"Per-Event Messaging","desc":"Host-configurable in-app messaging on mobile, visible to guests only when enabled.","stories":[
   {"id":"US-MSG-001","title":"Per-event messaging configuration","role":"host","want":"toggle guest messaging per event","benefit":"I control messaging per event","pri":"Medium","deps":"US-HOST-004",
    "ac":["An Allow guest messaging toggle is shown in the event's settings/messaging tab.","When off, guests see no messaging option.","Existing threads are preserved."]},
   {"id":"US-MSG-002","title":"Guest messaging gated by the toggle","role":"guest","want":"see Message host only when allowed","benefit":"messaging matches the host's setup","pri":"Medium","deps":"US-MSG-001",
    "ac":["Message host appears on the event/ticket only when enabled.","Disabled events hide the option entirely.","Threads open in-app."]},
   {"id":"US-MSG-003","title":"Staff messaging by permission","role":"staff member","want":"read/reply to messages only if permitted","benefit":"access matches my role","pri":"Low","deps":"US-STAFF-005",
    "ac":["Messaging screens appear only with the messaging permission.","Reply requires the reply permission.","Replies are attributed in the thread."]},
 ]},

 {"id":"PLAT","name":"Cross-cutting & Delivery","desc":"Non-functional concerns: responsive phone layout, the web build, deployment and prototype data.","stories":[
   {"id":"US-PLAT-001","title":"Aligned, responsive phone layout","role":"any user","want":"screens that are aligned and easy to use","benefit":"the app feels professional","pri":"High","deps":"—",
    "ac":["Content is consistently spaced and aligned; nothing is clipped under the status bar.","Lists, cards and images render correctly at phone widths.","Touch targets are usable."]},
   {"id":"US-PLAT-002","title":"Web build via react-native-web","role":"reviewer","want":"the app to run on the web inside a phone frame","benefit":"I can preview it in a browser","pri":"Medium","deps":"US-NAV-003",
    "ac":["The app builds for web (Expo export / react-native-web).","It renders inside a phone-frame device mockup.","Theming and behaviour match the handset."]},
   {"id":"US-PLAT-003","title":"Static web deployment","role":"product owner","want":"to host the web build for sharing","benefit":"stakeholders can try it via a link","pri":"Low","deps":"US-PLAT-002",
    "ac":["A build command produces a static web bundle.","Config exists for a static host (e.g. Netlify).","An SPA redirect keeps in-app navigation working."]},
   {"id":"US-PLAT-004","title":"Session persistence","role":"returning user","want":"my session remembered across visits (web)","benefit":"I skip re-logging in","pri":"Low","deps":"US-AUTH-005",
    "ac":["Guest/Host sessions persist on web between visits.","Staff sessions are not persisted (re-enter via Invite ID).","Log out clears the persisted session."]},
   {"id":"US-PLAT-005","title":"Prototype data store","role":"platform owner","want":"the app to run without a backend","benefit":"the prototype is self-contained","pri":"Medium","deps":"US-NAV-005",
    "ac":["All data is seeded in-app (no server).","A live store propagates state changes during a session.","Behaviour mirrors the web app's feature set."]},
 ]},
]

# ----------------------------------------------------------------------------
# DOCX BUILD HELPERS
# ----------------------------------------------------------------------------
def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hexcolor)
    tcPr.append(shd)

def set_cell_text(cell, text, bold=False, color=None, size=None, white=False):
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    if size: run.font.size = Pt(size)
    if white: run.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    elif color: run.font.color.rgb = color
    return p

def add_bullets(cell, items):
    cell.text = ''
    for i, it in enumerate(items):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.style = doc.styles['List Bullet']
        p.add_run(it)

doc = Document()
normal = doc.styles['Normal']
normal.font.name = 'Calibri'; normal.font.size = Pt(10.5); normal.font.color.rgb = INK
for h in ['Heading 1', 'Heading 2', 'Heading 3']:
    doc.styles[h].font.name = 'Calibri'; doc.styles[h].font.color.rgb = BRAND

# ---------- TITLE ----------
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('\n\n\n\nSafalEvents'); r.font.size = Pt(40); r.bold = True; r.font.color.rgb = BRAND
sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
rs = sub.add_run('Mobile App — End-to-End User Stories'); rs.font.size = Pt(19); rs.font.color.rgb = INK
sub2 = doc.add_paragraph(); sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
rs2 = sub2.add_run('Expo / React Native — Functional User Stories with Acceptance Criteria'); rs2.font.size = Pt(11); rs2.font.color.rgb = MUTED

doc.add_paragraph('\n\n')
info = doc.add_table(rows=5, cols=2); info.alignment = WD_TABLE_ALIGNMENT.CENTER; info.style = 'Table Grid'
info.columns[0].width = Inches(1.8); info.columns[1].width = Inches(3.6)
meta = [('Product', 'SafalEvents Mobile — Expo / React Native app (host, guest & staff)'),
        ('Document', 'Mobile End-to-End User Stories'),
        ('Version', '1.0'),
        ('Date', TODAY),
        ('Status', 'Draft for review')]
for i,(k,v) in enumerate(meta):
    set_cell_text(info.rows[i].cells[0], k, bold=True); shade(info.rows[i].cells[0], LABELBG)
    set_cell_text(info.rows[i].cells[1], v)
doc.add_page_break()

# ---------- HOW TO READ ----------
doc.add_heading('How to read this document', level=1)
doc.add_paragraph('This document captures the end-to-end functional scope of the SafalEvents mobile app as user stories, '
                  'grouped into epics (modules). Each story has a unique ID, a role-based narrative (As a … I want … so that …), '
                  'testable acceptance criteria, a priority and dependencies.')
conv = doc.add_table(rows=1, cols=2); conv.style = 'Table Grid'
hdr = conv.rows[0].cells
set_cell_text(hdr[0], 'Convention', bold=True, white=True); shade(hdr[0], HEADBG)
set_cell_text(hdr[1], 'Meaning', bold=True, white=True); shade(hdr[1], HEADBG)
for k,v in [('Story ID', 'US-<EPIC>-<NNN>, e.g. US-BROWSE-006. Stable references for traceability.'),
            ('Priority', 'High = core / must-have, Medium = important, Low = nice-to-have.'),
            ('Dependencies', 'Other story IDs that should be delivered first; "—" means none.'),
            ('Roles', 'Guest, Host, Staff (plus Reviewer / Platform Owner for cross-cutting items).')]:
    row = conv.add_row().cells
    set_cell_text(row[0], k, bold=True); shade(row[0], LABELBG)
    set_cell_text(row[1], v)

# ---------- EPIC SUMMARY ----------
doc.add_heading('Epic summary', level=1)
es = doc.add_table(rows=1, cols=4); es.style = 'Table Grid'
hc = es.rows[0].cells
for i,h in enumerate(['Epic ID','Epic','# Stories','Scope']):
    set_cell_text(hc[i], h, bold=True, white=True); shade(hc[i], HEADBG)
for ep in EPICS:
    c = es.add_row().cells
    set_cell_text(c[0], 'EPIC-'+ep['id'], bold=True)
    set_cell_text(c[1], ep['name'])
    set_cell_text(c[2], str(len(ep['stories'])))
    set_cell_text(c[3], ep['desc'])

# ---------- TRACEABILITY INDEX ----------
doc.add_heading('Traceability index (all stories)', level=1)
idx = doc.add_table(rows=1, cols=4); idx.style = 'Table Grid'
ic = idx.rows[0].cells
for i,h in enumerate(['Story ID','Title','Priority','Depends on']):
    set_cell_text(ic[i], h, bold=True, white=True); shade(ic[i], HEADBG)
total = 0
for ep in EPICS:
    for s in ep['stories']:
        total += 1
        c = idx.add_row().cells
        set_cell_text(c[0], s['id'], bold=True)
        set_cell_text(c[1], s['title'])
        set_cell_text(c[2], s['pri'])
        set_cell_text(c[3], s['deps'])
doc.add_page_break()

# ---------- DETAILED STORIES ----------
doc.add_heading('Detailed user stories', level=1)
for ep in EPICS:
    doc.add_heading('EPIC-%s — %s' % (ep['id'], ep['name']), level=2)
    doc.add_paragraph(ep['desc'])
    for s in ep['stories']:
        doc.add_heading('%s — %s' % (s['id'], s['title']), level=3)
        tbl = doc.add_table(rows=0, cols=2); tbl.style = 'Table Grid'
        tbl.columns[0].width = Inches(1.5); tbl.columns[1].width = Inches(5.0)
        def row(label, value=None, bullets=None):
            cells = tbl.add_row().cells
            set_cell_text(cells[0], label, bold=True); shade(cells[0], LABELBG)
            if bullets is not None:
                add_bullets(cells[1], bullets)
            else:
                set_cell_text(cells[1], value)
        row('Story ID', s['id'])
        row('Epic', 'EPIC-%s — %s' % (ep['id'], ep['name']))
        row('Priority', s['pri'])
        row('User story', 'As a %s, I want to %s, so that %s.' % (s['role'], s['want'], s['benefit']))
        row('Acceptance criteria', bullets=s['ac'])
        row('Dependencies', s['deps'])
        doc.add_paragraph('')

# ---------- FOOTER ----------
def add_page_number(footer_par):
    run = footer_par.add_run()
    fldStart = OxmlElement('w:fldChar'); fldStart.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve'); instr.text = 'PAGE'
    fldEnd = OxmlElement('w:fldChar'); fldEnd.set(qn('w:fldCharType'), 'end')
    run._r.append(fldStart); run._r.append(instr); run._r.append(fldEnd)

sec = doc.sections[0]
foot = sec.footer.paragraphs[0]; foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = foot.add_run('SafalEvents Mobile — End-to-End User Stories  ·  Page '); fr.font.size = Pt(8); fr.font.color.rgb = MUTED
add_page_number(foot)

if __name__ == "__main__":
    doc.save(OUT)
    print('Stories:', total)
    print('Saved:', OUT)
